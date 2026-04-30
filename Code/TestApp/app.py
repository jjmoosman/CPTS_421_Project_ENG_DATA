import sys
import re
import pandas as pd
import unicodedata
from pathlib import Path
from rapidfuzz import fuzz
from PyQt6.QtWidgets import (QApplication, QMainWindow, QPushButton, QVBoxLayout, 
                             QWidget, QTextEdit, QLineEdit, QFileDialog, QLabel, 
                             QProgressBar, QMessageBox)
from docx import Document
import fitz  # PyMuPDF


# A decorator that catches any exceptions in the decorated function and shows an error message box.
# This prevents the app from crashing and informs the user of problems.
def safe_slot(fn):
    def wrapper(self, *args, **kwargs):
        try:
            return fn(self, *args, **kwargs)
        except Exception as exc:
            QMessageBox.critical(self, "Application Error", f"Unexpected error: {exc}")
            return None
    return wrapper

class RedactorApp(QMainWindow):
    # Initializes the main application window, sets the title and size, and calls initUI to build the interface.
    # Also sets up empty lists and variables for storing user data.
    def __init__(self):
        super().__init__()
        self.setWindowTitle("FERPA Compliance Tool")
        self.setMinimumSize(600, 750)
        self.initUI()
        self.blacklisted_terms = set()
        self.selected_files = []
        self.output_dir = None
        self.temp_output_dir = None

    # Builds the user interface by creating buttons, text boxes, labels, and layouts.
    # Arranges everything in a vertical layout for the main window.
    def initUI(self):
        layout = QVBoxLayout()

        # Step 1: Input names and IDs
        layout.addWidget(QLabel("<b>Step 1: Input Names and IDs to Remove</b>"))
        
        prof_layout = QVBoxLayout()
        prof_layout.addWidget(QLabel("Professor Names:"))
        self.prof_input = QTextEdit()
        self.prof_input.setFixedHeight(80)
        prof_layout.addWidget(self.prof_input)
        layout.addLayout(prof_layout)
        
        name_layout = QVBoxLayout()
        name_layout.addWidget(QLabel("Student Names:"))
        self.name_input = QTextEdit()
        name_layout.addWidget(self.name_input)
        layout.addLayout(name_layout)
        
        id_layout = QVBoxLayout()
        id_layout.addWidget(QLabel("Student IDs:"))
        self.id_input = QTextEdit()
        id_layout.addWidget(self.id_input)
        layout.addLayout(id_layout)
        
        btn_excel = QPushButton("📁 Load from Excel")
        btn_excel.clicked.connect(self.load_excel)
        layout.addWidget(btn_excel)

        # Step 2: Custom file naming
        layout.addWidget(QLabel("<br><b>Step 2: Custom File Naming (Optional)</b>"))
        self.prefix_input = QLineEdit()
        self.prefix_input.setPlaceholderText("Enter file prefix (e.g. Class_Semester)")
        layout.addWidget(self.prefix_input)

        # Step 3: Select files
        layout.addWidget(QLabel("<br><b>Step 3: Select Student Files</b>"))
        file_btn_layout = QVBoxLayout()

        btn_files = QPushButton("📄 Select Individual Files")
        btn_files.clicked.connect(self.select_files)
        file_btn_layout.addWidget(btn_files)

        btn_folder = QPushButton("📂 Select Folder (All .docx/.pdf)")
        btn_folder.clicked.connect(self.select_folder)
        file_btn_layout.addWidget(btn_folder)
        layout.addLayout(file_btn_layout)

        btn_clear = QPushButton("🗑️ Clear Selection")
        btn_clear.clicked.connect(self.clear_selection)
        btn_clear.setStyleSheet("color: #d32f2f; font-size: 11px; border: 1px solid #d32f2f;")
        file_btn_layout.addWidget(btn_clear)

        self.file_list_label = QLabel("No files selected")
        layout.addWidget(self.file_list_label)

        # Step 4: Select destination
        layout.addWidget(QLabel("<br><b>Step 4: Select Destination Folder</b>"))
        btn_dest = QPushButton("Select Destination")
        btn_dest.clicked.connect(self.select_destination)
        layout.addWidget(btn_dest)

        self.selected_destination_label = QLabel("No destination selected")
        layout.addWidget(self.selected_destination_label)

        # De-identify button
        self.process_btn = QPushButton("🚀 DE-IDENTIFY")
        self.process_btn.clicked.connect(self.run_redaction)
        self.process_btn.setStyleSheet("background-color: #2E7D32; color: white; font-weight: bold; height: 40px;")
        layout.addWidget(self.process_btn)

        layout.addWidget(QLabel("<br><b>Results</b>"))
        self.results_output = QTextEdit()
        self.results_output.setReadOnly(True)
        self.results_output.setFixedHeight(220)
        self.results_output.setPlaceholderText("Processing results and logs will appear here. Scroll to view more output.")
        layout.addWidget(self.results_output)

        self.progress = QProgressBar()
        layout.addWidget(self.progress)
        
        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    # Opens a file dialog for the user to select an Excel file (.xlsx or .xls).
    # Reads the file, determines which columns contain names and IDs based on headers,
    # and populates the name and ID text boxes with the data.
    @safe_slot
    def load_excel(self, *args, **kwargs):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open Excel", "", "Excel Files (*.xlsx *.xls)")
        if file_path:
            path = Path(file_path)
            if not path.exists():
                QMessageBox.critical(self, "Excel Error", "Selected file does not exist.")
                return

            try:
                if path.suffix.lower() == ".xlsx":
                    df = pd.read_excel(path, engine="openpyxl")
                elif path.suffix.lower() == ".xls":
                    try:
                        df = pd.read_excel(path, engine="xlrd")
                    except ImportError:
                        QMessageBox.critical(self, "Excel Error", "To open .xls files install xlrd or use .xlsx files.")
                        return
                else:
                    raise ValueError("Unsupported Excel file type.")

                if df.empty:
                    QMessageBox.warning(self, "Excel Warning", "No data found in the selected Excel file.")
                    return

                # Determine columns for names and IDs
                headers = [str(col).lower() for col in df.columns]
                name_col = None
                id_col = None
                for i, header in enumerate(headers):
                    if 'name' in header:
                        name_col = i
                    elif 'id' in header:
                        id_col = i

                if name_col is None and len(df.columns) > 0:
                    name_col = 0  # Assume first column is names
                if id_col is None and len(df.columns) > 1:
                    id_col = 1  # Assume second column is IDs

                names = []
                ids = []

                for _, row in df.iterrows():
                    if name_col is not None:
                        name_val = str(row.iloc[name_col]).strip()
                        if name_val.lower() != 'nan' and name_val:
                            names.append(name_val)
                    if id_col is not None:
                        id_val = str(row.iloc[id_col]).strip()
                        if id_val.lower() != 'nan' and id_val:
                            ids.append(id_val)

                if names:
                    self.name_input.append("\n".join(names))
                if ids:
                    self.id_input.append("\n".join(ids))

                if not names and not ids:
                    QMessageBox.warning(self, "Excel Warning", "No valid data found in the selected Excel file.")
            except ImportError as e:
                QMessageBox.critical(self, "Excel Error", f"Missing package for Excel support: {e}")
            except Exception as e:
                QMessageBox.critical(self, "Excel Error", f"Failed to load Excel: {e}")

    # Opens a file dialog for the user to select individual .docx or .pdf files.
    # Adds the selected files to the list of files to process.
    @safe_slot
    def select_files(self, *args, **kwargs):
        files, _ = QFileDialog.getOpenFileNames(self, "Select Files", "", "Documents (*.docx *.pdf);;Word Documents (*.docx);;PDF Files (*.pdf)")
        if files:
            self.add_to_selected_files(files)

    # Opens a folder dialog for the user to select a directory.
    # Finds all .docx and .pdf files in that folder and adds them to the list.
    @safe_slot
    def select_folder(self, *args, **kwargs):
        folder_path = QFileDialog.getExistingDirectory(self, "Select Folder of Student Writings")
        if folder_path:
            folder = Path(folder_path)
            found_files = []
            for ext in ["*.pdf", "*.docx", "*.PDF", "*.DOCX"]:
                found_files.extend([str(f) for f in folder.glob(ext)])
            
            if found_files:
                self.add_to_selected_files(found_files)
            else:
                QMessageBox.warning(self, "No Files Found", "No PDF or DOCX files found in that folder.")

    # Adds new files to the selected files list without duplicates.
    # Updates the label showing the count and sets a default output directory if needed.
    def add_to_selected_files(self, new_files):
        combined = set(self.selected_files) | set(new_files)
        self.selected_files = list(combined)
        self.file_list_label.setText(f"Total files selected: {len(self.selected_files)}")
        
        if not self.output_dir and self.selected_files:
            self.temp_output_dir = Path(self.selected_files[0]).parent / "Redacted_Output"
            self.selected_destination_label.setText(f"Default: {self.temp_output_dir}")

    # Clears the list of selected files and resets the label.
    def clear_selection(self):
        self.selected_files = []
        self.file_list_label.setText("No files selected")

    def log_result(self, message: str):
        if hasattr(self, 'results_output'):
            self.results_output.append(message)
            self.results_output.verticalScrollBar().setValue(self.results_output.verticalScrollBar().maximum())
        else:
            print(message)
        
    # Parses pasted names/IDs into individual blacklist terms.
    # Supports newline, semicolon, tab separators and converts Last,First style names into full name terms.
    def parse_input_list(self, raw_text):
        entries = re.split(r'[\n;\t]+', raw_text)
        terms = []
        for entry in entries:
            entry = entry.strip()
            if not entry:
                continue

            if ',' in entry:
                parts = [p.strip() for p in entry.split(',') if p.strip()]
                if len(parts) == 2:
                    last, first = parts
                    terms.append(f"{first} {last}")
                    continue
                # Fall back to splitting across commas when there are more than two segments.
                for part in parts:
                    if part:
                        terms.append(part)
                continue

            terms.append(entry)

        return terms

    # Normalizes professor names by removing common titles like Dr., Prof., Mr., etc.
    def normalize_name(self, name):
        name = name.strip()
        name = re.sub(r'^(?:dr|prof|professor|mr|mrs|ms|miss|mx)\.?\s+', '', name, flags=re.IGNORECASE)
        name = re.sub(r'\s+', ' ', name).strip()
        return name

    # Opens a folder dialog for the user to choose where to save the redacted files.
    # Sets the output directory for saving results.
    @safe_slot
    def select_destination(self, *args, **kwargs):
        folder_path = QFileDialog.getExistingDirectory(self, "Select Destination Folder")
        if folder_path:
            selected_folder = Path(folder_path)
            self.output_dir = selected_folder / "Redacted_Output"
            self.selected_destination_label.setText(f"Destination: {self.output_dir}")

    # Starts the redaction process: collects names and IDs from text boxes,
    # checks for selected files, creates output directory, and processes each file.
    # Shows progress and a completion message.
    @safe_slot
    def run_redaction(self, *args, **kwargs):
        raw_professors = [self.normalize_name(t) for t in self.parse_input_list(self.prof_input.toPlainText())]
        raw_names = self.parse_input_list(self.name_input.toPlainText())
        raw_ids = self.parse_input_list(self.id_input.toPlainText())
        self.blacklisted_terms = {t for t in raw_professors + raw_names + raw_ids if len(t) > 1}
        header_terms = set(raw_professors + raw_names + raw_ids)
        for term in raw_professors + raw_names:
            if ' ' in term:
                parts = term.split()
                header_terms.update(parts)
        self.header_terms = {t for t in header_terms if len(t) > 1}
        
        if hasattr(self, 'results_output'):
            self.results_output.clear()
        self.log_result("Starting redaction...")
        self.selected_files = [f for f in self.selected_files if Path(f).exists()]
        if not self.selected_files or not self.blacklisted_terms:
            error_text = "Need files and names/IDs! Make sure selected files still exist."
            self.log_result(f"Error: {error_text}")
            QMessageBox.warning(self, "Missing Data", error_text)
            return
        
        # Determine final output path
        final_output = self.output_dir if self.output_dir else self.temp_output_dir
        if not final_output:
            QMessageBox.critical(self, "Error", "No output directory determined.")
            return

        try:
            final_output.mkdir(exist_ok=True, parents=True)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not create directory: {e}")
            return

        custom_prefix = self.prefix_input.text().strip()
        file_prefix = custom_prefix if custom_prefix else "Student"
        
        self.progress.setMaximum(len(self.selected_files))
        self.progress.setValue(0)

        for i, f_path in enumerate(self.selected_files):
            try:
                file_path = Path(f_path)
                out_path = final_output / f"{file_prefix}_{i+1}.txt"

                if file_path.suffix.lower() == ".docx":
                    self.redact_docx_to_txt(file_path, out_path)
                elif file_path.suffix.lower() == ".pdf":
                    self.redact_pdf_to_txt(file_path, out_path)
                else:
                    raise ValueError(f"Unsupported file type: {file_path.suffix}")
                self.log_result(f"Processed: {file_path.name} -> {out_path.name}")
            except Exception as e:
                error_msg = f"Could not process {f_path}: {e}"
                self.log_result(error_msg)
                QMessageBox.warning(self, "Processing Error", error_msg)
            
            self.progress.setValue(i + 1)
        
        finish_msg = f"Finished. Saved to: {final_output}"
        self.log_result(finish_msg)
        QMessageBox.information(self, "Finished", finish_msg)

    # Replaces occurrences of a term in the text with [REDACTED], using fuzzy matching for similarity.
    # Handles single words or full names by looking for matches in the text.
    def fuzzy_replace(self, text, term, threshold=80):
        term_lower = term.lower().strip()
        if not term_lower:
            return text

        redacted_text = text
        words = list(re.finditer(r'\b[\w-]+\b', text))

        if len(term_lower.split()) == 1:
            for i in range(len(words) - 1, -1, -1):
                word_match = words[i]
                word_clean = re.sub(r'[^a-zA-Z0-9]', '', word_match.group()).lower()
                if word_clean and fuzz.ratio(word_clean, term_lower) >= threshold:
                    redacted_text = redacted_text[:word_match.start()] + "[REDACTED]" + redacted_text[word_match.end():]

            pattern = rf'(?<!\w){re.escape(term_lower)}(?!\w)'
            redacted_text = re.sub(pattern, "[REDACTED]", redacted_text, flags=re.IGNORECASE)
            return redacted_text

        name_parts = [re.sub(r'[^a-zA-Z0-9]', '', part).lower() for part in term_lower.split() if part.strip()]
        if not name_parts:
            return redacted_text

        first, last = name_parts[0], name_parts[-1]
        first_initial = first[0] if first else ''
        honorific_prefix = r'(?:dr|prof|professor)\.?\s+'

        middle_token = r'(?:[A-Za-z]+\.?|[A-Za-z]\.)'
        separator = r'(?:[\s\.\-]+)'
        patterns = [
            rf'\b{re.escape(first)}[._-]{re.escape(last)}\b',
            rf'\b{re.escape(first_initial)}\.?\s+{re.escape(last)}\b',
            rf'\b{re.escape(last)}\s*,\s*{re.escape(first)}\b',
            rf'\b{re.escape(first)}(?:{separator}{middle_token}){{0,4}}{separator}{re.escape(last)}\b',
            rf'\b{honorific_prefix}{re.escape(first)}(?:{separator}{middle_token}){{0,4}}{separator}{re.escape(last)}\b'
        ]
        for p in patterns:
            redacted_text = re.sub(p, "[REDACTED]", redacted_text, flags=re.IGNORECASE)

        text_to_scan = redacted_text
        words = list(re.finditer(r'\b[\w-]+\b', text_to_scan))
        min_match = max(2, len(name_parts) - 1)
        match_ranges = []

        for seq_len in range(len(name_parts), min_match - 1, -1):
            for term_start in range(0, len(name_parts) - seq_len + 1):
                subseq = name_parts[term_start:term_start + seq_len]
                subseq_text = " ".join(subseq)
                for i in range(0, len(words) - seq_len + 1):
                    window = words[i:i + seq_len]
                    window_text = " ".join(
                        re.sub(r'[^a-zA-Z0-9]', '', w.group()).lower() for w in window
                    )
                    if fuzz.ratio(window_text, subseq_text) >= threshold:
                        match_ranges.append((window[0].start(), window[-1].end()))

        if match_ranges:
            match_ranges = sorted(match_ranges, key=lambda r: (r[0], -r[1]))
            merged_ranges = []
            for start, end in match_ranges:
                if not merged_ranges or start > merged_ranges[-1][1]:
                    merged_ranges.append([start, end])
                else:
                    merged_ranges[-1][1] = max(merged_ranges[-1][1], end)

            for start, end in reversed(merged_ranges):
                redacted_text = redacted_text[:start] + "[REDACTED]" + redacted_text[end:]

        return redacted_text

    # Applies redaction to the given text by replacing all blacklisted terms and emails with [REDACTED].
    # Processes terms in order of length to avoid conflicts.
    def redact_text(self, text):
        redacted_text = text
        for term in sorted(self.blacklisted_terms, key=lambda s: -len(s)):
            redacted_text = self.fuzzy_replace(redacted_text, term)
        redacted_text = self.redact_emails(redacted_text)
        return redacted_text

    # Finds email addresses in the text and redacts them if they contain any blacklisted terms.
    # Replaces matching emails with [REDACTED].
    def redact_emails(self, text):
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        redacted_text = text
        for email in emails:
            for term in self.blacklisted_terms:
                term_lower = term.lower().strip()
                if re.search(rf'(?<!\w){re.escape(term_lower)}(?!\w)', email.lower()):
                    redacted_text = redacted_text.replace(email, "[REDACTED]")
                    break
        return redacted_text

    # Yields all paragraphs from a DOCX document, including those in tables and headers/footers.
    # Helps to process text from all parts of the document.
    def iter_docx_paragraphs(self, doc):
        for para in doc.paragraphs:
            yield para
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para
        for section in doc.sections:
            for part in (section.header, section.footer):
                for para in part.paragraphs:
                    yield para
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                yield para

    # Extracts all text from a DOCX document by combining text from all paragraphs.
    # Returns the full text content as a single string.
    def get_docx_text(self, doc):
        return "\n".join(para.text for para in self.iter_docx_paragraphs(doc))

    # Processes a DOCX file: extracts text, redacts it, saves as .txt, and verifies no terms remain.
    def redact_docx_to_txt(self, input_path, output_path):
        try:
            doc = Document(str(input_path))
        except Exception as e:
            raise RuntimeError(f"Unable to open DOCX file: {e}")

        raw_text = self.get_docx_text(doc)
        redacted_text = self.redact_text(raw_text)
        self.write_text_output(output_path, redacted_text)
        self.verify_text_output(output_path)

    # Extracts text from a single PDF page using PyMuPDF.
    def extract_pdf_page_text(self, page):
        return page.get_text("text")

    # Cleans up corrupted special characters found in poorly encoded PDFs.
    # Converts ligatures, accented characters, and smart quotes to their ASCII equivalents.
    # This ensures clean, readable text in the output file.
    def clean_pdf_text(self, text):
        text = unicodedata.normalize('NFKC', text)

        # Fix common PDF encoding artifacts
        text = re.sub(r'Ɵ', 'ti', text)

        replacements = {
            '–': '-',
            '—': '-',
            '“': '"',
            '”': '"',
            '‘': "'",
            '’': "'",
        }

        for k, v in replacements.items():
            text = text.replace(k, v)

        return text

    # Processes a PDF file: extracts text from all pages, redacts it, saves as .txt, and verifies.
    def redact_pdf_to_txt(self, input_path, output_path):
        try:
            doc = fitz.open(str(input_path))
        except Exception as e:
            raise RuntimeError(f"Unable to open PDF file: {e}")

        try:
            page_texts = []
            for page in doc:
                page_text = self.extract_pdf_page_text(page)
                page_text = self.clean_pdf_text(page_text)
                lines = page_text.split('\n')
                if lines and any(term.lower() in lines[0].lower() for term in self.header_terms):
                    lines[0] = "[REDACTED HEADER]"
                page_texts.append('\n'.join(lines))
            raw_text = "\n".join(page_texts)
            redacted_text = self.redact_text(raw_text)
            self.write_text_output(output_path, redacted_text)
            self.verify_text_output(output_path)
        except Exception as e:
            raise RuntimeError(f"Failed to extract or verify redacted PDF text: {e}")
        finally:
            doc.close()

    # Writes the given text to a file at the specified path.
    def write_text_output(self, output_path, text):
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
        except Exception as e:
            raise RuntimeError(f"Failed to write text output: {e}")

    # Checks the output file to ensure no blacklisted terms remain.
    # Raises an error if any are found.
    def verify_text_output(self, output_path):
        try:
            with open(output_path, "r", encoding="utf-8") as f:
                content = f.read().lower()
        except Exception as e:
            raise RuntimeError(f"Unable to verify text output: {e}")

        found = [term for term in self.blacklisted_terms if term.lower() in content]
        if found:
            raise RuntimeError(f"Verification failed. Unredacted terms found in text output: {', '.join(sorted(set(found)))}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = RedactorApp()
    window.show()
    sys.exit(app.exec())